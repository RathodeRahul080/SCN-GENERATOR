
from flask import Flask, request, send_file
from docx import Document
from docx.shared import Inches
import io
import datetime

app = Flask(__name__)

ISSUE_TEMPLATES = {
    "Non-payment": "It was observed that the taxpayer has failed to discharge the GST liability for the declared tax period. The failure to pay tax has resulted in loss of revenue and contravenes Section 9(1) of the CGST Act, 2017.",
    "Wrong ITC": "It was observed that the taxpayer has wrongly claimed input tax credit without actual receipt of goods/services or in excess of eligible credit. This constitutes misstatement and contravenes Section 74 of the CGST Act, 2017.",
    "Non-filing": "It was found that the taxpayer has not filed GST returns for the specified tax period. This non-compliance violates the provisions of Section 39 and 44 of the CGST Act, 2017.",
    "Others": "Irregularities were found in taxpayer’s GST filings which require explanation under relevant provisions of the GST law."
}

@app.route('/generate', methods=['POST'])
def generate_doc():
    data = request.json
    doc = Document()
    doc.add_heading('SHOW CAUSE NOTICE', 0)

    doc.add_paragraph(f"OC No.: {data['oc_number']}      Date: {data['date_of_notice']}")
    doc.add_paragraph(f"DIN: {data['din']}")

    doc.add_heading(f"Subject: GST – {data['subject']} – Issuance of Show Cause Notice – Reg.", level=1)

    doc.add_paragraph(f"M/s. {data['taxpayer_name']}, located at {data['taxpayer_address']} (GSTIN: {data['taxpayer_gstin']}), is under scrutiny for the following matter:")
    doc.add_paragraph(ISSUE_TEMPLATES.get(data['issue_type'], data['issue_description']))
    doc.add_paragraph(f"Tax Period: {data['tax_period']}")

    doc.add_heading("Summary of Tax Liability", level=2)
    doc.add_paragraph(f"IGST: ₹{data['igst']}")
    doc.add_paragraph(f"CGST: ₹{data['cgst']}")
    doc.add_paragraph(f"SGST: ₹{data['sgst']}")
    doc.add_paragraph(f"CESS: ₹{data['cess']}")
    doc.add_paragraph(f"Total Liability: ₹{data['total_tax']}")

    doc.add_heading("Legal Grounds and Explanation", level=2)

    doc.add_paragraph("6. The taxpayer has failed to comply with the GST law by not fulfilling their tax obligations. This is considered a serious contravention as it directly impacts the revenue stream. Such omissions indicate wilful default and are not in accordance with the principles of self-assessment.")

    doc.add_paragraph("7. Under Section 74(1) of the CGST Act, 2017, the liability to pay unpaid tax along with interest is enforceable. As per the provided records, the taxpayer has not responded to notices and hence the determination of liability stands justified.")

    doc.add_paragraph("8. The taxpayer appears to have violated the following: Section 9(1) for failure to pay tax, Section 50(1) for non-payment of interest, and other related rules under CGST and APGST.")

    doc.add_paragraph("9. The legal foundation for taking action is further supported by corresponding provisions under the APGST Act and IGST Act. These laws operate in harmony and provide a legal basis for the present SCN.")

    doc.add_paragraph("10. Both CGST and APGST Acts are parallel in most respects. For practical implementation, references to CGST shall also imply references to APGST unless otherwise stated.")

    doc.add_paragraph(f"11. DRC-01A (Ref: {data['drco_reference']}) dated {data['drco_date']} was served but no clarification was received.")

    doc.add_paragraph(f"12. Therefore, M/s {data['taxpayer_name']} is hereby required to show cause within 30 days from the date of this notice as to why the tax amount of ₹{data['total_tax']} should not be recovered along with interest and penalty.")

    doc.add_paragraph("13. The taxpayer is also directed to furnish all documentary evidence they intend to rely upon in their reply and indicate if they wish to be heard in person.")

    doc.add_paragraph("14. Failure to respond will result in ex-parte adjudication of this matter based on the available records.")

    doc.add_paragraph("15. This notice is without prejudice to any other legal action under GST law or any other law currently in force.")

    doc.add_paragraph("16. The notice is based on verified documents such as invoices, returns and e-way bills already available with the department.")

    doc.add_paragraph("\n\nSUPERINTENDENT\nNANDYAL-I CGST RANGE")
    doc.add_paragraph(f"To: {data['taxpayer_name']}, {data['taxpayer_address']}\nEmail: {data['taxpayer_email']}")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(file_stream, as_attachment=True, download_name='Show_Cause_Notice.docx')

if __name__ == '__main__':
    app.run(debug=True)
